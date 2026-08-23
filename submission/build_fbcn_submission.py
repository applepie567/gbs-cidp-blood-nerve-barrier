from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re, csv, shutil

ROOT=Path('/workspace/scratch/70ab84f1d47d')
SRC=ROOT/'output/Revised_GBS_CIDP_BNB_manuscript_humanized.docx'
OUT=ROOT/'output/FBCNS_submission'
OUT.mkdir(parents=True,exist_ok=True)

TITLE='Cross-platform multi-omics analysis maps a myeloid–blood–nerve barrier–Schwann cell circuit across Guillain–Barré syndrome and CIDP'
RUNNING='A blood–nerve barrier circuit in GBS and CIDP'
AUTHORS='Yuan Feng¹, Fang Feng², Shuxiao Ma¹, Shuman Feng¹, Chao Jiang¹*, Yibin Hao¹*'
AFF1='¹ [Department, institution, city, postal code, country — AUTHOR CONFIRMATION REQUIRED]'
AFF2='² [Department, institution, city, postal code, country — AUTHOR CONFIRMATION REQUIRED]'
CORR='* Correspondence: Chao Jiang ([email]) and Yibin Hao ([email]). Full postal address and telephone — AUTHOR CONFIRMATION REQUIRED.'

src=Document(SRC)
paras=[p.text.strip() for p in src.paragraphs]

def get_refs():
    start=paras.index('References')+1
    end=paras.index('Figure legends')
    refs={}
    for t in paras[start:end]:
        m=re.match(r'^(\d+)\.\s*(.*)$',t)
        if m: refs[int(m.group(1))]=m.group(2)
    return refs
REFS=get_refs()

abstract=[
('Background','The blood–nerve barrier (BNB) governs immune access to peripheral nerve, but the relationship between circulating inflammation, BNB endothelial responses and target-organ remodeling across Guillain–Barré syndrome (GBS) and chronic inflammatory demyelinating polyradiculoneuropathy (CIDP) remains incompletely defined.'),
('Methods','We integrated four acute GBS blood transcriptomic cohorts, paired acute-to-one-year serum proteomics, a human CIDP sural-nerve single-nucleus atlas and a normal human BNB reference. Prespecified programs captured CXCL8–CXCR1/2 recruitment, complement, Fc receptors, interferon/JAK–STAT signaling, transendothelial migration, BNB identity, macrophage states and Schwann-cell repair. Blood effects were synthesized by random-effects meta-analysis with Hartung–Knapp inference. Single-nucleus counts were aggregated by donor and cell state. Differential programs, composition and same-donor ligand–receptor–target coordination were evaluated without treating cells as independent replicates.'),
('Results','Acute GBS converged on a circulating myeloid recruitment program involving CXCL8–CXCR1/2, inflammatory monocytes, complement, Fc receptors and transendothelial migration. CXCL8–CXCR1/2 was positive in all four blood cohorts (random-effects Hedges g=1.18, I²=0%, Hartung–Knapp P=0.053). Paired proteomics separated declining interferon-related activity from persistent LIF/OSM and myeloid signals. CIDP nerve showed cell-state-selective remodeling. Complement and Fc-receptor programs localized to macrophages. Venular and capillary endothelial states retained junctional and adhesion machinery while showing altered interferon and antigen-presentation programs. Repair Schwann cells showed altered IRF1 and antigen-presentation activity. Same-donor analysis revealed coordinated changes across macrophage, endothelial and Schwann-cell programs. However, no individual interaction score remained significant after correction for multiple testing.'),
('Conclusions','The data map a BNB-centered sequence in which acute circulating recruitment is linked to chronic target-organ remodeling. This myeloid–BNB–Schwann cell framework prioritizes experimentally testable chemokine, adhesion, complement, Fc-receptor and trophic pathways for longitudinal studies of inflammatory peripheral neuropathy.')]

background=[
'Guillain–Barré syndrome (GBS) is an acute immune-mediated polyradiculoneuropathy that usually reaches clinical nadir within four weeks, whereas chronic inflammatory demyelinating polyradiculoneuropathy (CIDP) follows a progressive or relapsing course [1–3]. Despite their different time scales, both disorders involve leukocyte recruitment, humoral immunity, macrophage-mediated myelin injury and Schwann-cell responses. Acute-onset CIDP and treatment-related fluctuations further show that early clinical boundaries do not always identify the underlying trajectory.',
'Immune entry into peripheral nerve is regulated by the blood–nerve barrier (BNB), formed by tight-junction-bearing endoneurial microvascular endothelial cells and associated pericytes [4–7]. The BNB is anatomically and transcriptionally distinct from the blood–brain barrier and from the nonvascular perineurium. Endoneurial endothelial ICAM1, VCAM1, CCL2 and leukocyte integrins participate in inflammatory adhesion and transendothelial migration [7–9]. These features position the BNB as the anatomical interface linking circulating immune recruitment to peripheral nerve injury.',
'Beyond entry, antibody and complement recognition engage macrophage Fc receptors and phagocytic programs, whereas injury-responsive Schwann cells adopt repair states that support debris clearance and axonal regeneration [20,21,32–44]. These observations suggest a multicellular sequence—circulating recruitment, endothelial engagement, macrophage effector activity and Schwann-cell repair—but the components have largely been examined in separate cohorts and compartments.',
'We therefore integrated acute GBS blood transcriptomics, paired longitudinal serum proteomics, donor-resolved CIDP nerve single-nucleus data and a normal human BNB reference. We investigated which circulating programs were reproducible across acute GBS cohorts, how macrophage, BNB endothelial and Schwann-cell states were organized in CIDP nerve, and whether same-donor expression supported candidate routes linking these compartments. The analysis was designed to map expression-supported relationships and phase-specific programs rather than infer direct causal signaling.']

methods_sections=[
('Study design and datasets',[paras[16],paras[17]]),
('Definition and scoring of biological programs',['Before performing group comparisons, we defined gene programs representing key components of immune-mediated peripheral nerve injury. These programs covered CXCL8–CXCR1/2-dependent leukocyte recruitment, LIF/OSM signaling through LIFR, OSMR and IL6ST, complement activation and regulation, Fc-gamma receptor signaling, interferon–JAK–STAT activity, inflammatory-monocyte responses, endothelial adhesion and transendothelial migration, BNB junctional integrity, macrophage residency and phagocytosis, and Schwann-cell myelin, stress and repair responses. Each program was scored using genes detected in the corresponding dataset. This approach enabled comparison of the same biological processes across platforms while retaining dataset-specific measurements.']),
('Acute GBS cohort analysis and meta-analysis',[paras[21]]),
('Longitudinal proteomic analysis',[paras[23]]),
('CIDP donor-by-cell-state pseudobulk',[paras[25]]),
('Pathway, regulon and cellular composition analysis',[paras[27]]),
('Same-donor intercellular analysis',[paras[29].replace('Communication analysis','Intercellular analysis')]),
('Statistical analysis',[paras[31]])]

results_sections=[
('Study architecture centers the BNB within a multi-compartment design',[paras[34].replace('(Fig. 1)','(Fig. 1 and Table 1)')]),
('Acute GBS converges on a circulating myeloid recruitment program',[
    'GSE211225 showed a broad acute response in whole blood. Compared with healthy controls, patients with acute GBS had higher activity of the CXCL8–CXCR1/2 program (Hedges g=1.25, FDR=0.045), LIF/OSM–gp130 signaling (Hedges g=2.19, FDR=0.0076), complement (Hedges g=2.49, FDR=0.0076), Fc-receptor signaling (Hedges g=1.36, FDR=0.038), interferon–JAK–STAT signaling (Hedges g=1.07, FDR=0.045) and the inflammatory-monocyte program (Hedges g=1.93, FDR=0.015). This pattern linked chemotactic recruitment to humoral effector activity rather than to a single dominant transcript (Fig. 2).',
    'Comparison with post-acute samples identified the programs most closely associated with the acute phase. Acute GBS showed higher complement activity (Hedges g=1.35, FDR=0.041), Fc-receptor signaling (Hedges g=1.23, FDR=0.046), transendothelial migration (Hedges g=1.59, FDR=0.024) and inflammatory-monocyte activity (Hedges g=1.68, FDR=0.024). CXCL8–CXCR1/2 and LIF/OSM–gp130 signaling remained directionally elevated but showed smaller phase contrasts. These findings distinguished an acute recruitment and effector program from responses that may persist beyond neurological nadir.',
    'Independent datasets localized much of this response to the myeloid compartment. In GSE31014, the effects for CXCL8–CXCR1/2, Fc-receptor signaling and transendothelial migration were positive, and the inflammatory-monocyte program had a Hedges g of 1.17. Sorted CD11b+ cells from patients with early untreated AIDP showed positive effects for CXCL8, OSM, C3, CCL2 and CCR1. In PRJNA1293757, CXCL8 activity was concentrated in monocytes. Although the cohorts differed in platform and cellular composition, each associated chemokine-driven recruitment with the acute myeloid response.',
    'All seven prespecified programs had positive random-effects estimates across the four acute GBS datasets. CXCL8–CXCR1/2 was positive in every cohort and showed no detectable between-study heterogeneity (random-effects Hedges g=1.18, I²=0%, Hartung–Knapp P=0.053). The summary Hedges g was 1.01 for the inflammatory-monocyte program and 0.65 for transendothelial migration. Complement had a summary Hedges g of 0.76 with moderate heterogeneity (I²=56.5%). Effects for LIF/OSM–gp130 and Fc-receptor signaling varied more widely among whole blood, sorted cells and PBMC monocytes (Fig. 2).',
    'Leave-one-cohort-out analysis preserved the positive direction of the CXCL8–CXCR1/2 program. Removing individual cohorts widened the confidence intervals but did not reverse the effect. At the gene level, 89 transcripts were measured in at least three datasets, although their effects were less stable than the program-level estimates. The cross-cohort analysis therefore identified recruitment as a coordinated biological feature and CXCL8–CXCR1/2 as its most reproducible component.'
]),
('Longitudinal proteomics resolves distinct recovery trajectories',[
    'The paired proteomic cohort distinguished persistent circulating differences from changes occurring during recovery. The LIF/OSM–gp130 and inflammatory-monocyte protein programs were higher in acute GBS than in healthy controls and remained elevated at one year. Interferon–JAK–STAT activity declined within participants (paired P=0.0168, FDR=0.118). CXCL8–CXCR1/2 and transendothelial-migration programs showed smaller decreases that varied according to the statistical test. Complement and LIF/OSM–gp130 signaling did not show a uniform decline from the acute phase to recovery (Fig. 3).',
    'Transcriptomic and proteomic effects captured different aspects of the response. Among 65 genes represented in both acute blood transcriptomics and serum proteomics, the Spearman correlation was 0.031 and directional concordance was 58.5%. Cellular transcription, secretion, tissue release and protein clearance therefore provided distinct molecular views of the same clinical phase. Candidates supported by cellular localization and pathway coherence were more informative than those selected solely on the basis of RNA–protein agreement.'
]),
('CIDP nerve resolves macrophage, BNB endothelial and Schwann-cell programs',[
    'The sural-nerve atlas resolved the proposed circuit within the target organ. Complement and Fc-receptor genes were concentrated in macrophage states. Endothelial states expressed CLDN5, OCLN, TJP1 and CDH5 together with IL6ST, LIFR, OSMR, ICAM1 and VCAM1. Schwann-cell states expressed gp130-family receptors alongside genes involved in myelin maintenance, cellular stress and repair. Granulocytes showed the strongest CXCR1/2 expression, whereas endothelial expression was lower. This distribution positioned CXCL8 primarily as a signal for leukocyte recruitment and myeloid amplification (Fig. 4).',
    'Normal endoneurial microvessels in GSE107574 showed a similar vascular framework, including CLDN5, OCLN, TJP1, CDH5, VWF, LIFR, IL6ST, ICAM1 and C3. The human BNB reference therefore linked the endothelial states identified in CIDP nerve to a specialized peripheral vascular interface equipped for cytokine reception, complement regulation and leukocyte adhesion.',
    'Donor-level analysis across all expressed genes did not identify a single transcript that passed state-wide FDR correction. Pathway analysis nevertheless revealed a coordinated shift in Macro2 macrophages. SPI1 target activity (FDR=0.0022), the resident-macrophage program (FDR=0.0027), Fc-gamma phagocytosis (FDR=0.0046) and interferon/IRF7 activity (FDR=0.0046) were lower in CIDP than in CIAP. This pattern was consistent with the lower macrophage Fc-receptor score observed in CIDP (difference in standardized score=-0.673, FDR=0.0087).',
    paras[50],
    'Endothelial changes were state specific. CIDP was associated with lower interferon-response and IRF7 programs in the first venular/capillary endothelial subtype (both FDR=0.0035). Antigen-presentation activity was lower in the second venular/capillary endothelial subtype (FDR=0.0037). Junctional and adhesion genes remained readily detectable, indicating preservation of endothelial identity alongside altered immune-response programs. Together, these endothelial populations define a vascular setting for leukocyte arrest and passage into endoneurial tissue.',
    paras[53],
    'Compositional analysis identified a lower epithelial-like cell fraction in CIDP after centered-log-ratio transformation (FDR=0.037). Damage Schwann cells, nonmyelinating Schwann cells and the second venular/capillary endothelial subtype showed directional abundance changes that did not pass FDR correction. The dominant CIDP signal was therefore transcriptional reorganization within cell states, accompanied by a restricted change in cellular composition.'
]),
('Same-donor coupling supports a BNB-centered tissue circuit',[paras[56].replace('(Fig. 5)','(Fig. 5)'),paras[57],paras[58]]),
('Cross-compartment synthesis links recruitment to target-organ remodeling',[paras[60].replace('(Fig. 6)','(Fig. 6 and Tables 2 and 3)')])]

discussion=[
'This cross-platform analysis places the BNB at the center of a multicellular response spanning acute GBS blood and chronically injured CIDP nerve. Acute GBS was dominated by a reproducible myeloid recruitment program, whereas CIDP tissue contained selective macrophage, endothelial and Schwann-cell remodeling. The results therefore support a change in the location and organization of immune activity rather than a single inflammatory scale shared by both diseases.',
'CXCL8–CXCR1/2 was the most directionally stable acute program. Whole blood showed simultaneous chemokine, complement, Fc-receptor, interferon and inflammatory-monocyte activity, and independent leukocyte, CD11b+ and PBMC datasets localized much of this response to myeloid populations. Its consistency across heterogeneous sources is compatible with elevated IL-8 in GBS biofluids and with chemokine- and integrin-dependent leukocyte trafficking at human endoneurial endothelium [7,20,32,53].',
'The cellular distribution clarifies how CXCL8 may operate at the BNB. CXCR1/2 expression was strongest in granulocytic and myeloid populations, whereas venular/capillary endothelial states expressed junctional and adhesion machinery. CXCL8 is therefore positioned primarily as a luminal recruitment and myeloid-amplification signal rather than a dominant endothelial receptor pathway. A direct test would expose human BNB models to patient biofluids under flow and quantify leukocyte arrest, passage and permeability with CXCR1/2 or ICAM1 blockade.',
'Paired proteomics showed that recovery was not a uniform reversal of the acute transcriptome. Interferon-related proteins declined most clearly within participants, whereas LIF/OSM and inflammatory-myeloid programs remained elevated. LIF-family signaling can support Schwann-cell survival and repair [42–44], so persistence may reflect continuing tissue restoration as well as residual inflammation. The weak RNA–protein correlation also indicates that cellular transcription, secretion, tissue release and clearance provide different views of disease phase.',
'CIDP nerve was characterized by cell-state-selective remodeling. Macro2 macrophages showed coordinated differences in SPI1, resident-macrophage, Fc-phagocytosis and interferon programs relative to CIAP. Because CIAP is itself an injured-nerve comparator, these differences are more consistent with altered state selection, chronic immune exposure and treatment history than with a simple increase in macrophage abundance. IVIg-associated changes in activating and inhibitory Fc-receptor balance provide one plausible contribution [21], but treatment effects cannot be isolated in the available tissue cohort.',
'BNB endothelial and repair Schwann-cell states shared altered interferon and antigen-presentation programs while retaining vascular or glial identity. This pattern places endoneurial endothelium between systemic recruitment and intraneural responses and suggests that antigen handling is coordinated across multiple resident compartments. The normal BNB reference independently confirmed expression of junctional, cytokine-receptor, complement-regulatory and adhesion genes in human endoneurial microvessels.',
'Same-donor analysis provided expression-level support for the proposed tissue circuit, including inflammatory macrophage ligands associated with endothelial adhesion and Schwann-cell response modules. These results do not establish direct cell–cell signaling: diagnosis-associated interaction scores were modest, and spatial proximity was not measured. Spatial transcriptomics, multiplex imaging and perturbation in human BNB co-culture systems will be needed to determine whether the predicted source and receiver populations occupy and communicate within the same endoneurial niches.',
'The study is limited by small blood cohorts, platform and tissue heterogeneity, and the absence of matched blood, cerebrospinal fluid and nerve from the same participants. CIDP was compared primarily with CIAP rather than healthy nerve, and disease duration and treatment may contribute to tissue-state differences. These constraints are partly addressed by participant/donor-level inference, random-effects synthesis and explicit separation of direct measurements from expression-supported mechanisms. A prospective design should include pretreatment acute GBS, recovery, active CIDP and timed pre/post-IVIg sampling combined with BNB flow assays and cell-resolved protein measurements.']

conclusion='Across blood, serum and peripheral nerve, inflammatory neuropathy was organized around a myeloid–BNB–Schwann cell axis. Acute GBS showed reproducible CXCL8-centered recruitment, whereas CIDP nerve displayed selective macrophage Fc/interferon remodeling, BNB endothelial antigen-response changes and Schwann-cell repair programs. The resulting barrier-centered model provides specific hypotheses for longitudinal biomarkers and functional studies of leukocyte trafficking at the human BNB.'

legends=[
('Figure 1','Study architecture and BNB-centered analytical framework. Four acute GBS blood cohorts, paired longitudinal serum proteomics, CIDP sural-nerve single-nucleus data and a normal human BNB reference were organized around recruitment, barrier engagement, macrophage effector activity and Schwann-cell repair. BNB, blood–nerve barrier, denotes the specialized endoneurial microvascular endothelial barrier of peripheral nerve and is not used as a synonym for the blood–brain barrier.'),
('Figure 2','Acute GBS converges on myeloid recruitment. (A) Hedges g for seven prespecified programs across whole blood, leukocytes, sorted CD11b+ cells and PBMC monocytes. (B) Random-effects estimates with 95% confidence intervals; teal points indicate nominal P<0.05. (C) Whole-blood acute-versus-control and acute-versus-post-acute contrasts. (D) Sample-level monocyte scores in PBMC data. Participants, not cells, were the inferential unit.'),
('Figure 3','Longitudinal proteomics resolves recovery trajectories. (A) Paired acute-to-one-year patient trajectories for selected protein programs; orange and teal points denote acute and one-year measurements, respectively, and lines connect measurements from the same participant. (B) Within-participant standardized effects; teal denotes paired P<0.05 and navy denotes P≥0.05. Displayed P values are paired-test results. (C) Cross-sectional acute-versus-control and one-year-versus-control effects. (D) Biological interpretation separates recruitment, transient interferon response and persistent remodeling.'),
('Figure 4','CIDP nerve resolves barrier and repair programs. (A) Expression localization of endothelial, macrophage and Schwann-cell genes across nerve states. Dot size represents detection and color represents scaled mean expression. (B) CIDP-versus-CIAP program shifts; diamonds denote FDR<0.05. (C) Donor-level differences in centered-log-ratio cell fractions. (D) Spearman associations between program scores and INCAT disability; orange denotes nominal P<0.05 and teal denotes P≥0.05. BNB EC denotes endoneurial microvascular endothelial cells.'),
('Figure 5','Cell-state coupling supports a BNB-centered tissue circuit. (A) Cross-compartment ligand and receptor evidence linking acute GBS recruitment to CIDP nerve receiver states. (B) Same-donor standardized programs illustrate heterogeneity among CIDP nerves. (C) Expression-supported model connecting circulating myeloid cells, BNB endothelium, nerve macrophages and Schwann-cell injury/repair. Same-donor expression supports a coordinated tissue circuit, while the magnitude of individual programs varies among donors. These relationships are hypothesis-generating and do not demonstrate direct signaling.'),
('Figure 6','Integrated barrier-centered model across disease phases. Acute GBS emphasizes circulating myeloid recruitment; recovery separates transient from persistent protein programs; active chronic nerve contains macrophage, BNB endothelial and Schwann-cell remodeling; and stable or treated CIDP may retain treatment-shaped Fc and repair programs. Here, BNB refers specifically to the blood–nerve barrier formed by specialized endoneurial microvascular endothelium in peripheral nerve. The diagram represents a phase-aware mechanistic synthesis rather than a longitudinal observation of the same patients.')]

def parse_cites(text):
    nums=[]
    for m in re.finditer(r'\[(\d+(?:\s*[–-]\s*\d+)?(?:\s*,\s*\d+(?:\s*[–-]\s*\d+)?)*)\]',text):
        for part in re.split(r'\s*,\s*',m.group(1)):
            if re.search(r'[–-]',part):
                a,b=map(int,re.split(r'\s*[–-]\s*',part)); nums.extend(range(a,b+1))
            else: nums.append(int(part))
    return nums

all_body=[]
all_body += background
for _,xs in methods_sections: all_body+=xs
for _,xs in results_sections: all_body+=xs
all_body += discussion+[conclusion]
order=[]
for t in all_body:
    for n in parse_cites(t):
        if n in REFS and n not in order: order.append(n)
for n in sorted(REFS):
    if n not in order: order.append(n)
mapping={old:i+1 for i,old in enumerate(order)}

def compress(ns):
    ns=sorted(dict.fromkeys(ns)); out=[]; i=0
    while i<len(ns):
        j=i
        while j+1<len(ns) and ns[j+1]==ns[j]+1: j+=1
        out.append(f'{ns[i]}–{ns[j]}' if j-i>=2 else ','.join(str(x) for x in ns[i:j+1]))
        i=j+1
    return ','.join(out)
def remap(text):
    def f(m):
        old=[]
        for part in re.split(r'\s*,\s*',m.group(1)):
            if re.search(r'[–-]',part):
                a,b=map(int,re.split(r'\s*[–-]\s*',part)); old.extend(range(a,b+1))
            else: old.append(int(part))
        return '['+compress([mapping[x] for x in old if x in mapping])+']'
    return re.sub(r'\[(\d+(?:\s*[–-]\s*\d+)?(?:\s*,\s*\d+(?:\s*[–-]\s*\d+)?)*)\]',f,text)

def setup(doc, line_spacing=2):
    sec=doc.sections[0]; sec.top_margin=Inches(.8); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.85); sec.right_margin=Inches(.75)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(10); normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')
    normal.paragraph_format.line_spacing=line_spacing; normal.paragraph_format.space_after=Pt(0 if line_spacing==2 else 5)
    for s,size in [('Heading 1',13),('Heading 2',11)]:
        st=styles[s]; st.font.name='Arial'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor(0,0,0); st._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')
        st.paragraph_format.keep_with_next=True; st.paragraph_format.space_before=Pt(8); st.paragraph_format.space_after=Pt(2); st.paragraph_format.line_spacing=1
    # continuous line numbering
    sectPr=sec._sectPr; ln=OxmlElement('w:lnNumType'); ln.set(qn('w:countBy'),'1'); ln.set(qn('w:restart'),'continuous'); ln.set(qn('w:distance'),'360'); sectPr.append(ln)
    # footer page number
    p=sec.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); p._p.append(fld)

def add_title(doc, full=True):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(TITLE); r.bold=True; r.font.name='Arial'; r.font.size=Pt(16)
    p.paragraph_format.line_spacing=1.15; p.paragraph_format.space_after=Pt(8)
    for t,b in [(AUTHORS,True),(AFF1,False),(AFF2,False),(CORR,False)]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(t); rr.bold=b; rr.font.size=Pt(10); p.paragraph_format.line_spacing=1.15
    if full:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Article type: Research Article').italic=True
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Running title: '+RUNNING).italic=True

def add_h(doc,t,level=1): doc.add_paragraph(t,style=f'Heading {level}')
def add_p(doc,t,boldlabel=None):
    t=remap(str(t)).replace(';', ',')
    p=doc.add_paragraph()
    if boldlabel:
        r=p.add_run(boldlabel+': '); r.bold=True; p.add_run(t)
    else: p.add_run(t)
    return p

def add_table_rows(doc, rows, title, widths=None):
    add_p(doc,title); doc.paragraphs[-1].runs[0].bold=True
    tab=doc.add_table(rows=1,cols=len(rows[0])); tab.alignment=WD_TABLE_ALIGNMENT.CENTER; tab.autofit=False; tab.style='Table Grid'
    for j,v in enumerate(rows[0]): tab.rows[0].cells[j].text=v.replace('_',' ').replace(';', ',')
    for r in rows[1:]:
        cells=tab.add_row().cells
        for j,v in enumerate(r): cells[j].text=str(v).replace(';', ',')
    for ri,row in enumerate(tab.rows):
        for j,c in enumerate(row.cells):
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths: c.width=Inches(widths[j])
            for p in c.paragraphs:
                p.paragraph_format.line_spacing=1; p.paragraph_format.space_after=Pt(0)
                for run in p.runs: run.font.name='Arial'; run.font.size=Pt(7.5); run.bold=(ri==0)
    add_p(doc,'Values are derived from participant- or donor-level analyses. Blank cells indicate unavailable or non-estimable quantities.')
    doc.paragraphs[-1].runs[0].italic=True; doc.paragraphs[-1].paragraph_format.line_spacing=1

def add_table_from_csv(doc, filename, title, keep_cols=None, widths=None, decimals=None):
    rows=list(csv.reader(open(ROOT/'output/figure_source_data'/filename,encoding='utf-8')))
    if keep_cols:
        rows=[[r[i] if i<len(r) else '' for i in keep_cols] for r in rows]
    if decimals:
        for r in rows[1:]:
            for j,d in decimals.items():
                if j<len(r) and r[j] not in ('','NA','NaN'):
                    try: r[j]=f'{float(r[j]):.{d}f}'
                    except ValueError: pass
    add_table_rows(doc,rows,title,widths)

def build_manuscript():
    d=Document(); setup(d); add_title(d)
    add_h(d,'Abstract')
    for label,text in abstract: add_p(d,text,label)
    add_p(d,'Guillain–Barré syndrome, chronic inflammatory demyelinating polyradiculoneuropathy, blood–nerve barrier, endoneurial endothelial cell, macrophage, Schwann cell, CXCL8, complement','Keywords')
    add_h(d,'Background')
    for t in background: add_p(d,t)
    add_h(d,'Methods')
    for h,xs in methods_sections:
        add_h(d,h,2)
        for t in xs: add_p(d,t)
    add_h(d,'Results')
    for h,xs in results_sections:
        add_h(d,h,2)
        for t in xs: add_p(d,t)
    add_h(d,'Discussion')
    for t in discussion: add_p(d,t)
    add_h(d,'Conclusions'); add_p(d,conclusion)
    add_h(d,'List of abbreviations'); add_p(d,paras[75])
    add_h(d,'Declarations')
    declarations=[
      ('Ethics approval and consent to participate','This study reanalyzed publicly available, de-identified human datasets and did not recruit participants or obtain new specimens. The original studies reported ethics approval and informed consent. No additional institutional review-board approval was required for the present secondary analysis. This statement must be confirmed against the authors’ institutional policy before submission.'),
      ('Consent for publication','Not applicable.'),
      ('Availability of data and materials','All primary datasets are publicly available through GEO, NCBI BioProject or the source repositories under the accession numbers listed in the Methods and References. Source data underlying the figures and tables are included in Additional file 1 and in the accompanying public GitHub repository. No new individual-level clinical dataset was generated.'),
      ('Code availability','Versioned analysis scripts, figure-generation code and aggregate source data are available at https://github.com/applepie567/gbs-cidp-blood-nerve-barrier. A permanent DOI will be added after the verified release is archived in Zenodo.'),
      ('Competing interests','[AUTHOR CONFIRMATION REQUIRED: if none, replace with “The authors declare that they have no competing interests.”]'),
      ('Funding','[AUTHOR CONFIRMATION REQUIRED: provide funder names, grant numbers and the role of each funder, or state that the study received no specific funding.]'),
      ('Authors’ contributions','YF conceived the study, developed the analytical framework, performed the analyses and drafted the manuscript. FF curated data, validated results and revised the manuscript. SXM contributed to validation, interpretation and manuscript revision. SMF contributed to data curation, literature evaluation and manuscript revision. CJ supervised the methodology and critically revised the manuscript. YH supervised the study, coordinated the project and critically revised the manuscript. All authors must confirm that they have read and approved the final manuscript.'),
      ('Acknowledgements','We thank the participants and investigators who generated and shared the public datasets used in this study. During preparation of the manuscript, the authors used OpenAI ChatGPT Work and Codex to assist with code scaffolding, documentation and language editing. The authors reviewed and verified all outputs and take full responsibility for the content of the manuscript.'),
      ('Authors’ information','Not applicable.')]
    for h,t in declarations: add_h(d,h,2); add_p(d,t)
    add_h(d,'Additional files')
    for t in [
      'Additional file 1 (.xlsx): Machine-readable figure, table and panel-level source data.',
      'Additional file 2 (.docx): Reproducibility appendix, dataset provenance, software environment and extended methods.',
      'Additional file 3 (.zip): Versioned analysis code and configuration files.']:
        add_p(d,t)
    add_h(d,'Tables')
    dataset_rows=[['Resource','Material','Participants or samples','Analytical role'],
      ['GSE211225','Whole-blood transcriptomics','6 acute GBS; 10 post-acute; 6 controls','Phase contrasts and program effects'],
      ['GSE31014','Blood leukocyte microarray','7 GBS; 7 controls','Independent direction check'],
      ['GSE304871','Sorted blood RNA-seq','CD11b+: 2/3; CD4+: 3/3; CD8+: 2/3 case/control','Early untreated immune-source localization'],
      ['PRJNA1293757','PBMC single-cell RNA-seq','3 untreated AIDP; 2 controls','Sample-level monocyte localization'],
      ['GBS-Proteomics','Serum SomaScan','20 acute/recovery pairs; 15 controls','Within-patient one-year trajectories'],
      ['GSE285983','Sural-nerve single-nucleus RNA-seq','9 CIDP; 11 CIAP; 37 total donors','Donor-by-cell-state target-organ analysis'],
      ['GSE107574','Human endoneurial endothelial RNA','2 cultured preparations; 4 microvessel preparations','Descriptive normal BNB identity reference']]
    add_table_rows(d,dataset_rows,'Table 1. Datasets and analytical roles',widths=[1.25,1.7,2.3,2.3])
    add_table_from_csv(d,'Table_3_crosscohort_robustness.csv','Table 2. Cross-cohort acute GBS program synthesis',keep_cols=[0,1,2,3,4,5,6],widths=[1.5,.85,.75,.75,.9,.75,.65],decimals={1:3,2:3,3:3,4:3,5:1,6:0})
    add_table_from_csv(d,'Table_2_cross_compartment_axes.csv','Table 3. Cross-compartment signaling axes',keep_cols=[0,1,2,3,4],widths=[1.4,1.3,1.1,1.1,2.3],decimals={2:3,3:3})
    add_h(d,'References')
    for new,old in enumerate(order,1): add_p(d,f'{new}. {REFS[old]}')
    add_h(d,'Figure legends')
    for label,text in legends: add_p(d,text,label)
    p=OUT/'Manuscript_FBCNS_final.docx'; d.save(p); return p

def build_titlepage():
    d=Document(); setup(d,1.15); add_title(d)
    add_h(d,'Author details')
    add_p(d,'ORCID identifiers, academic degrees and individual email addresses for all authors — AUTHOR CONFIRMATION REQUIRED.')
    add_h(d,'Corresponding authors')
    add_p(d,'Chao Jiang: [degree], [institution], [postal address], [email], [telephone], [ORCID].')
    add_p(d,'Yibin Hao: [degree], [institution], [postal address], [email], [telephone], [ORCID].')
    add_h(d,'Author note'); add_p(d,'Confirm whether any authors contributed equally and add the appropriate symbol only after all authors approve.')
    p=OUT/'Title_page_FBCNS.docx'; d.save(p); return p

def build_cover():
    d=Document(); setup(d,1.15)
    add_p(d,'22 August 2026'); add_p(d,'Editors-in-Chief\nFluids and Barriers of the CNS')
    add_p(d,'Dear Editors,')
    texts=[
      f'We submit the Research Article entitled “{TITLE}” for consideration in Fluids and Barriers of the CNS.',
      'The study addresses a barrier system explicitly covered by the journal: the blood–nerve barrier. By integrating four acute GBS blood transcriptomic cohorts, paired longitudinal serum proteomics, donor-resolved CIDP nerve single-nucleus data and a normal human BNB reference, we connect circulating immune recruitment with endoneurial endothelial, macrophage and Schwann-cell programs.',
      'Across acute GBS cohorts, CXCL8–CXCR1/2 was the most directionally stable recruitment program, accompanied by inflammatory-monocyte, complement, Fc-receptor and transendothelial-migration signals. In CIDP nerve, donor-level analyses localized complementary components of the proposed circuit to venular/capillary endothelial, macrophage and repair Schwann-cell states. Same-donor expression supports coordinated tissue programs, while the manuscript clearly distinguishes these hypotheses from direct or causal cell–cell signaling.',
      'The work is relevant to the journal because it treats peripheral endoneurial microvascular endothelium as the BNB—not the blood–brain barrier—and focuses on immune surveillance, leukocyte passage, barrier response and repair within human peripheral nerve. Panel-level source data and versioned analysis scripts accompany the submission in machine-readable form.',
      'All authors must confirm before submission that they have approved the manuscript, that the work has not been published or submitted elsewhere, and that all competing interests and funding sources have been fully declared. The study used public de-identified datasets and involved no new participant recruitment or specimen collection.',
      'Thank you for considering this manuscript.' ]
    for t in texts: add_p(d,t)
    add_p(d,'Sincerely,\n\nChao Jiang and Yibin Hao\nCorresponding authors\n[Affiliation, postal address, email and telephone — AUTHOR CONFIRMATION REQUIRED]')
    p=OUT/'Cover_letter_FBCNS.docx'; d.save(p); return p

def build_checklist():
    d=Document(); setup(d,1.0)
    p=d.add_paragraph(); r=p.add_run('Fluids and Barriers of the CNS: pre-submission checklist'); r.bold=True; r.font.size=Pt(16)
    add_p(d,'Manuscript: '+TITLE)
    items=[
      ('Journal fit','BNB is the organizing interface; BBB and perineurium are defined separately.','Ready'),
      ('Article type','Research Article selected.','Ready'),('Abstract','Structured Background/Methods/Results/Conclusions; below 350 words; no citations.','Ready'),
      ('Main text','Background, Methods, Results, Discussion and Conclusions; verify final main-text count is below 6000 words.','Check after author edits'),
      ('Formatting','DOCX, double-line spacing, continuous line numbering and page numbering; no manual page breaks.','Ready'),
      ('Authors','Verify spelling, order, degrees, ORCIDs, affiliations and all email addresses.','Required'),
      ('Correspondence','Complete postal addresses, emails and telephone numbers for Chao Jiang and Yibin Hao.','Required'),
      ('Ethics','Confirm institutional position on secondary analysis of public de-identified data.','Required'),
      ('Consent for publication','Not applicable unless individual-level identifiable material is added.','Ready'),
      ('Competing interests','Replace placeholder with the approved declaration for all authors.','Required'),
      ('Funding','Add every funder, grant number and funder role, or state no specific funding.','Required'),
      ('Contributions','All six authors must approve the contribution statement and final manuscript.','Required'),
      ('Data availability','Primary accessions listed; machine-readable derived data designated Additional file 1.','Ready'),
      ('Code availability','Analysis code designated Additional file 3; add GitHub URL and Zenodo DOI if deposited.','Required if claiming public archive'),
      ('AI disclosure','Acknowledgements states the limited use of generative AI and confirms author verification and responsibility.','Ready'),
      ('References','Vancouver numbering reordered by first appearance; verify preprints and 2026 records immediately before submission.','Final check'),
      ('Figures','Six composite figures; titles ≤15 words; legends ≤300 words; keys are inside figures.','Ready'),
      ('Figure files','Upload separate vector PDFs or 600-dpi PNGs; each file <10 MB; verify fonts and 170-mm legibility.','Ready'),
      ('Tables','Three editable Word tables; no color/shading; numbered in citation order.','Ready'),
      ('Additional files','Use exact names Additional file 1–3 and cite each in the manuscript.','Ready'),
      ('Cover letter','Add final corresponding-author signature block; confirm exclusivity, author approval and competing interests.','Required'),
      ('Submission proof','Inspect converted PDF, figure order, symbols, line numbers, supplementary labels and hyperlinks.','Required'),
      ('Final authority','Corresponding author completes declarations and clicks Submit after coauthor approval.','Required')]
    tab=d.add_table(rows=1,cols=3); tab.style='Table Grid'; tab.alignment=WD_TABLE_ALIGNMENT.CENTER; tab.autofit=False
    for j,v in enumerate(['Item','Requirement','Status']): tab.rows[0].cells[j].text=v
    for a,b,c in items:
        row=tab.add_row().cells
        for j,v in enumerate([a,b,c]): row[j].text=v.replace(';', ',')
    for ri,row in enumerate(tab.rows):
        for j,c in enumerate(row.cells):
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            c.width=Inches([1.35,5.2,1.15][j])
            for p in c.paragraphs:
                p.paragraph_format.line_spacing=1; p.paragraph_format.space_after=Pt(0)
                for run in p.runs: run.font.size=Pt(8); run.bold=(ri==0)
    add_h(d,'Files to upload')
    for t in ['Main manuscript DOCX','Title page DOCX if requested separately by the portal','Cover letter DOCX','Figures 1–6 as separate files','Additional files 1–3']: add_p(d,t)
    add_h(d,'Current journal requirements checked')
    add_p(d,'Fluids and Barriers of the CNS Research submission guidelines accessed 22 August 2026: structured abstract ≤350 words; main text ≤6000 words; double spacing; line and page numbering; mandatory Declarations headings; figure titles ≤15 words and legends ≤300 words; machine-readable data encouraged.')
    p=OUT/'Pre_submission_checklist_FBCNS.docx'; d.save(p); return p

files=[build_manuscript(),build_titlepage(),build_cover(),build_checklist()]
print('\n'.join(str(x) for x in files))
